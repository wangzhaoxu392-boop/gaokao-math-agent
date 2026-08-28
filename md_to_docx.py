# -*- coding: utf-8 -*-
"""将《开发历程记录.md》转换为同格式的 .docx（python-docx）"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "开发历程记录.md"
DST = "开发历程记录.docx"

def add_runs_with_bold(para, text):
    """解析 **加粗** 与 `代码` 并写入段落"""
    # 先按 **bold** 切
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.bold = True
        else:
            # 再处理 `code`
            sub = re.split(r"(`[^`]*`)", part)
            for s in sub:
                if s.startswith("`") and s.endswith("`") and len(s) > 1:
                    r = para.add_run(s[1:-1])
                    r.font.name = "Consolas"
                    r.font.size = Pt(10)
                elif s:
                    para.add_run(s)

def main():
    with open(SRC, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    # 中文字体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue
        # 分隔线
        if stripped in ("---", "***"):
            doc.add_paragraph("―" * 20)
            i += 1
            continue
        # 引用
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            r = p.add_run(stripped.lstrip("> ").strip())
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            r.italic = True
            i += 1
            continue
        # 标题
        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            p = doc.add_heading(level=level)
            add_runs_with_bold(p, text)
            i += 1
            continue
        # 表格
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|[\s:\-|]+\|\s*$", lines[i+1]):
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            # 解析
            def parse_row(r):
                cells = r.strip().strip("|").split("|")
                return [c.strip() for c in cells]
            header = parse_row(table_lines[0])
            body = [parse_row(r) for r in table_lines[2:]]
            rows = len(body) + 1
            cols = max(len(header), max((len(r) for r in body), default=0))
            table = doc.add_table(rows=rows, cols=cols)
            table.style = "Table Grid"
            for j, h in enumerate(header):
                cell = table.cell(0, j)
                cell.text = ""
                add_runs_with_bold(cell.paragraphs[0], h)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for ri, row in enumerate(body, start=1):
                for j in range(cols):
                    val = row[j] if j < len(row) else ""
                    cell = table.cell(ri, j)
                    cell.text = ""
                    add_runs_with_bold(cell.paragraphs[0], val)
            continue
        # 列表项（- 或 1. 2. 等）
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, m.group(2))
            i += 1
            continue
        m = re.match(r"^(\s*)(\d+)[.、]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs_with_bold(p, f"{m.group(2)}. {m.group(3)}")
            i += 1
            continue
        # 普通段落
        p = doc.add_paragraph()
        add_runs_with_bold(p, stripped)
        i += 1

    doc.save(DST)
    print("已生成", DST)

if __name__ == "__main__":
    main()
