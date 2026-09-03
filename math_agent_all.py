from dotenv import load_dotenv
import os
import re
import json
import time
import sqlite3
import requests
from pathlib import Path
import matplotlib.pyplot as plt
import sympy
from langchain_ollama import ChatOllama, OllamaEmbeddings
from langgraph.graph import MessagesState, StateGraph
from langgraph.prebuilt import ToolNode
from langchain_core.tools import tool
from langgraph.checkpoint.sqlite import SqliteSaver
from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
import pymupdf as fitz  # PyMuPDF，PDF文本提取+OCR
from docx import Document

plt.switch_backend('Agg')
load_dotenv()

# ========================全局路径配置========================
CHROMA_DB_PATH = "./db_chroma"
KNOWLEDGE_FOLDER = "./knowledge_docs"
BATCH_INPUT_FOLDER = "./batch_input"
OUTPUT_FOLDER = "./output"
RAW_EXAM_FOLDER = "./raw_exam_files"
OUT_RESEARCH_FOLDER = "./output_research"
# Karpathy LLM Wiki：编译产物与检索库（功能2 新架构）
WIKI_FOLDER = "./wiki"
WIKI_DB_PATH = "./db_wiki"

for p in [CHROMA_DB_PATH, KNOWLEDGE_FOLDER, BATCH_INPUT_FOLDER, OUTPUT_FOLDER, RAW_EXAM_FOLDER, OUT_RESEARCH_FOLDER, WIKI_FOLDER, WIKI_DB_PATH]:
    os.makedirs(p, exist_ok=True)

# 扩展 LangGraph 状态：增加模型选择字段，支持按题型自动路由 / 手动切换
class SolveState(MessagesState):
    model_choice: str   # "auto" / "reasoning" / "math"
    model_used: str     # 实际使用的模型名称（用于输出标注）

llm = ChatOllama(
    base_url=os.getenv("OLLAMA_BASE_URL"),
    model=os.getenv("LLM_MODEL"),
    temperature=0.2,
    num_predict=3072,   # 单次输出上限，防止 deepseek-r1 无限思考/超长输出
    timeout=900         # 15 分钟超时兜底，防止单次推理永久卡死
)

# 快速模型：用于拆题等轻量环节，加快功能1速度（可用 .env 的 FAST_MODEL 覆盖）
llm_fast = ChatOllama(
    base_url=os.getenv("OLLAMA_BASE_URL"),
    model=os.getenv("FAST_MODEL", "qwen2.5:7b"),
    temperature=0.1,
    num_predict=2048,
    timeout=300
)

# 数学专用模型：支持本地 Ollama（免费）或硅基流动 API（免费额度内，72B数学模型更准）
# .env 配置：MATH_PROVIDER=ollama 或 siliconflow；MATH_MODEL=模型名
_MATH_PROVIDER = os.getenv("MATH_PROVIDER", "ollama").lower()
_MATH_MODEL_NAME = os.getenv("MATH_MODEL", "qwen2.5:14b")

if _MATH_PROVIDER == "siliconflow":
    from langchain_openai import ChatOpenAI
    llm_math = ChatOpenAI(
        model=_MATH_MODEL_NAME,
        api_key=os.getenv("SILICONFLOW_API_KEY", ""),
        base_url="https://api.siliconflow.cn/v1",
        temperature=0.1,
        max_tokens=3072,
        timeout=120
    )
    print(f"[模型] 数学模型使用硅基流动 API: {_MATH_MODEL_NAME}")
else:
    llm_math = ChatOllama(
        base_url=os.getenv("OLLAMA_BASE_URL"),
        model=_MATH_MODEL_NAME,
        temperature=0.1,
        num_predict=3072,
        timeout=900
    )
    print(f"[模型] 数学模型使用本地 Ollama: {_MATH_MODEL_NAME}")

CATEGORY_LIST = ["集合", "函数", "导数", "三角", "数列", "立体几何", "圆锥曲线", "概率统计"]

# ========================API余额查询与用量统计========================
_USAGE_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "api_usage.json")

def _load_usage():
    if os.path.exists(_USAGE_FILE):
        try:
            with open(_USAGE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {"api_problems_solved": 0, "estimated_cost_cny": 0.0, "last_balance": None}

def _save_usage(data):
    try:
        with open(_USAGE_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def record_api_usage(estimated_cost=0.015):
    """记录一次 API 数学模型解题用量"""
    if _MATH_PROVIDER != "siliconflow":
        return
    data = _load_usage()
    data["api_problems_solved"] = data.get("api_problems_solved", 0) + 1
    data["estimated_cost_cny"] = round(data.get("estimated_cost_cny", 0.0) + estimated_cost, 4)
    _save_usage(data)

def query_siliconflow_balance():
    """查询硅基流动账户余额。返回 (成功标志, 消息文本)"""
    if _MATH_PROVIDER != "siliconflow":
        return False, f"当前数学模型使用本地 Ollama（{_MATH_MODEL_NAME}），不涉及 API 费用。"
    api_key = os.getenv("SILICONFLOW_API_KEY", "")
    if not api_key:
        return False, "未配置 SILICONFLOW_API_KEY，请在 .env 中设置。"
    try:
        resp = requests.get(
            "https://api.siliconflow.com/v1/user/info",
            headers={"Authorization": f"Bearer {api_key}"},
            timeout=8
        )
        if resp.status_code == 200:
            d = resp.json().get("data", {})
            usage = _load_usage()
            msg = (
                f"💰 硅基流动账户余额\n"
                f"  可用余额：¥{d.get('balance', '?')}\n"
                f"  充值余额：¥{d.get('chargeBalance', '?')}\n"
                f"  总余额：¥{d.get('totalBalance', '?')}\n"
                f"  账户状态：{d.get('status', '?')}\n"
                f"\n📊 本地用量统计\n"
                f"  API 解题次数：{usage.get('api_problems_solved', 0)}\n"
                f"  预估费用：¥{usage.get('estimated_cost_cny', 0):.4f}\n"
                f"  （每题约 ¥0.01~0.02，推理题走本地模型免费）"
            )
            usage["last_balance"] = d.get("balance")
            _save_usage(usage)
            return True, msg
        else:
            usage = _load_usage()
            msg = (
                f"⚠️ 余额查询失败（HTTP {resp.status_code}）\n"
                f"  硅基流动国内站 API Key 可能不支持在线查询余额。\n"
                f"  请手动查看：https://siliconflow.cn/account/fee\n\n"
                f"📊 本地用量统计\n"
                f"  API 解题次数：{usage.get('api_problems_solved', 0)}\n"
                f"  预估费用：¥{usage.get('estimated_cost_cny', 0):.4f}\n"
                f"  剩余额度 ≈ 初始余额 - 预估费用（请在费用中心确认准确余额）"
            )
            return False, msg
    except Exception as e:
        usage = _load_usage()
        msg = (
            f"⚠️ 余额查询异常：{str(e)[:100]}\n"
            f"  请手动查看：https://siliconflow.cn/account/fee\n\n"
            f"📊 本地用量统计\n"
            f"  API 解题次数：{usage.get('api_problems_solved', 0)}\n"
            f"  预估费用：¥{usage.get('estimated_cost_cny', 0):.4f}"
        )
        return False, msg

# ========================工具定义========================
@tool
def sympy_math_calc(code: str) -> str:
    """符号数学计算：求导、解方程、不等式、数列求和。
    已预定义符号 x,y,z,a,b,c,n,k,t（均为 sympy.Symbol），可直接使用；也可自行定义新符号。
    解方程时自动给出数值解并标注实根/复根。"""
    try:
        namespace = {"sympy": sympy, "sp": sympy}
        for name in ["x", "y", "z", "a", "b", "c", "n", "k", "t"]:
            namespace[name] = sympy.Symbol(name)
        # 先尝试 eval（纯表达式），失败则 exec（含赋值语句，取 result/ans 变量）
        try:
            res = eval(code, namespace)
        except SyntaxError:
            exec(code, namespace)
            res = namespace.get("result", namespace.get("ans", "（执行完成，无返回值）"))
        # 若是解列表（解方程结果），自动转数值并标注实根/复根，避免返回复杂精确式
        if isinstance(res, (list, tuple)):
            lines = [f"计算结果（共{len(res)}个解）："]
            for i, r in enumerate(res):
                try:
                    num = sympy.N(r, 6)
                    is_real = bool(sympy.im(r).simplify() == 0)
                    tag = "实根" if is_real else "复根"
                    exact_str = str(r)
                    if len(exact_str) > 80:
                        exact_str = exact_str[:77] + "..."
                    lines.append(f"  解{i+1} [{tag}]: 数值={num}, 精确={exact_str}")
                except Exception:
                    lines.append(f"  解{i+1}: {r}")
            return "\n".join(lines)
        # 普通结果：若表达式复杂，同时给数值近似
        try:
            num = sympy.N(res, 6)
            if str(num) != str(res) and len(str(res)) > 40:
                return f"计算结果：{res}\n数值近似：{num}"
        except Exception:
            pass
        return f"计算结果：{res}"
    except Exception as e:
        return f"计算工具错误：{str(e)}"

@tool
def plot_function(expr_str: str, x_range: tuple = (-10, 10)) -> str:
    """绘制函数图像，保存图片到output文件夹"""
    x = sympy.Symbol('x')
    expr = sympy.parse_expr(expr_str)
    f_lambda = sympy.lambdify(x, expr, "matplotlib")
    import numpy as np
    xs = np.linspace(x_range[0], x_range[1], 500)
    ys = f_lambda(xs)
    plt.figure(figsize=(6, 4))
    plt.plot(xs, ys)
    plt.grid(True)
    plt.title(f"y = {expr_str}")
    img_path = os.path.join(OUTPUT_FOLDER, "plot_out.png")
    plt.savefig(img_path, dpi=120)
    plt.close()
    return f"图像已保存 {img_path}"

tools = [sympy_math_calc, plot_function]
llm_with_tools = llm.bind_tools(tools)
llm_math_with_tools = llm_math.bind_tools(tools)

# ======================== 模型路由：按题型自动选择推理模型 / 数学模型 ========================
_REASONING_KEYWORDS = re.compile(
    r"证明|求证|是否存在|存不存在|探究|探讨|讨论|说明理由|为什么|分析.*原因|"
    r"比较.*大小|判断.*是否|试问|阐述|论述"
)

def route_model(question: str) -> str:
    """根据题目关键词自动路由：返回 'reasoning'（DeepSeek推理）或 'math'（Qwen数学）。
    证明/探究/存在性等推理密集型 → DeepSeek-R1；
    计算/方程/圆锥曲线/概率等计算密集型（默认）→ Qwen2.5-Math。"""
    if _REASONING_KEYWORDS.search(question):
        return "reasoning"
    return "math"

def resolve_model(choice: str, question: str):
    """根据用户选择（auto/reasoning/math）+ 题目，返回 (带工具模型, 纯文本模型, 模型显示名)。"""
    if choice == "reasoning":
        return llm_with_tools, llm, "DeepSeek-R1（推理模型）"
    if choice == "math":
        return llm_math_with_tools, llm_math, f"数学模型（{_MATH_PROVIDER}:{_MATH_MODEL_NAME}）"
    # auto
    routed = route_model(question)
    if routed == "reasoning":
        return llm_with_tools, llm, "DeepSeek-R1（推理模型·自动）"
    return llm_math_with_tools, llm_math, f"数学模型（{_MATH_PROVIDER}:{_MATH_MODEL_NAME}·自动）"

def get_llm_for_choice(choice: str, default: str = "llm"):
    """根据 model_choice 返回 (LLM实例, 显示名)，用于直接调用 .invoke() 的场景（非图模式）。
    choice: auto / reasoning / math
    default: auto 时用哪个模型 —— 'llm_fast'(快速编译) / 'llm'(推理分析) / 'llm_math'(数学计算)"""
    if choice == "reasoning":
        return llm, "DeepSeek-R1（推理模型）"
    if choice == "math":
        return llm_math, f"数学模型（{_MATH_PROVIDER}:{_MATH_MODEL_NAME}）"
    # auto
    if default == "llm_fast":
        return llm_fast, "Qwen2.5-7B（快速模型·自动）"
    if default == "llm_math":
        return llm_math, f"数学模型（{_MATH_PROVIDER}:{_MATH_MODEL_NAME}·自动）"
    return llm, "DeepSeek-R1（推理模型·自动）"

# ========================B RAG知识库模块【升级支持pdf、docx】========================
def read_knowledge_file(filepath: Path) -> str:
    """读取知识库文件：txt md docx pdf"""
    suffix = filepath.suffix.lower()
    if suffix == ".pdf":
        text_buf = []
        doc = fitz.open(filepath)
        for page in doc:
            pt = page.get_text("text")
            if pt:
                text_buf.append(pt)
        doc.close()
        return "\n".join(text_buf)
    elif suffix == ".docx":
        doc = Document(filepath)
        return "\n".join([p.text for p in doc.paragraphs])
    elif suffix in (".txt", ".md"):
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    else:
        return ""

def build_rag_vector_db():
    all_text = []
    support_exts = (".txt", ".md", ".docx", ".pdf")
    for fname in os.listdir(KNOWLEDGE_FOLDER):
        fp = Path(KNOWLEDGE_FOLDER) / fname
        if fp.suffix.lower() in support_exts:
            print(f"正在读取知识库文件：{fname}")
            content = read_knowledge_file(fp)
            if content.strip():
                all_text.append(content)
    if not all_text:
        print("knowledge_docs文件夹没有有效文档，跳过构建向量库")
        return None
    full_doc = "\n".join(all_text)
    splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=80)
    chunks = splitter.split_text(full_doc)
    embeddings = OllamaEmbeddings(model="nomic-embed-text")
    db = Chroma.from_texts(chunks, embedding=embeddings, persist_directory=CHROMA_DB_PATH)
    print("✅RAG向量知识库构建完成，支持格式：txt/md/docx/pdf")
    return db

def load_rag_db():
    embeddings = OllamaEmbeddings(model="nomic-embed-text")
    if os.path.exists(CHROMA_DB_PATH) and len(os.listdir(CHROMA_DB_PATH)) > 0:
        db = Chroma(persist_directory=CHROMA_DB_PATH, embedding_function=embeddings)
        return db
    return None

def rag_retrieve(query: str):
    # 【Karpathy 改造】优先从编译好的 LLM Wiki 检索结构化词条；未编译/失败则回退旧 RAG
    wiki_ctx = wiki_retrieve(query)
    if wiki_ctx:
        return wiki_ctx
    # 旧 RAG 回退：若 knowledge_docs 无任何可检索文档，直接跳过向量库检索，
    # 避免每次解题都加载嵌入模型（nomic-embed-text）拖慢速度。
    doc_files = [
        f for f in Path(KNOWLEDGE_FOLDER).glob("*")
        if f.is_file() and f.suffix.lower() in (".txt", ".md", ".docx", ".pdf")
    ]
    if not doc_files:
        return "【RAG】无知识库文档"
    db = load_rag_db()
    if db is None:
        return "【RAG】无知识库文档"
    docs = db.similarity_search(query, k=3)
    context = "\n====参考知识点片段====\n"
    for d in docs:
        context += d.page_content + "\n"
    return context

# ========================B2 Karpathy LLM Wiki 模块（功能2 新架构）========================
# 思想：不再每次从原始文档临时检索（RAG），而是让 LLM 把 knowledge_docs 资料
# “编译”成结构化的 Markdown Wiki（词条+交叉链接+总索引），提问时直接读 Wiki。
# 层次：knowledge_docs/（原始资料层）→ wiki/（编译词条层）→ 编译规范（内置 prompt）。
WIKI_COMPILE_PROMPT = """
你是高中数学教研专家。下面是一份讲义/知识点资料片段，请把它“编译”成 Wiki 词条。
规则：
1. 提炼出资料中真正的核心概念/知识点作为词条标题，标题简短有力（如“导数与单调性”“等比数列求和”“离心率”）。不要用一句话原文当标题。
2. 每个词条字段：
   - title：词条标题
   - category：分类，只能从【集合、函数、导数、三角、数列、立体几何、圆锥曲线、概率统计、其他】选一个
   - summary：一句话概述该知识点
   - key_points：要点数组，逐条列出关键结论/公式/方法
   - related：相关联的其它词条标题数组（用于交叉链接）
3. 只输出 JSON 数组，不要任何多余解释或 Markdown 代码块标记。
资料片段：
{chunk}
"""

def _split_knowledge_text(text: str, size: int = 800):
    """先按空行切段，再合并成约 size 字符的块；单段过长则硬切，保证每块不超过 size，
    避免整份大文件挤成一块导致模型 JSON 输出过长被截断而解析失败。"""
    parts = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks, cur = [], ""
    for p in parts:
        # 单段超过 size：先清空当前缓冲，再硬切成长段
        while len(p) > size:
            if cur:
                chunks.append(cur)
                cur = ""
            chunks.append(p[:size])
            p = p[size:]
        if cur and len(cur) + len(p) > size:
            chunks.append(cur)
            cur = p
        else:
            cur = (cur + "\n" + p).strip()
    if cur:
        chunks.append(cur)
    return chunks or ([text.strip()] if text.strip() else [])


def _safe_wiki_name(title: str) -> str:
    """词条标题转合法文件名"""
    t = re.sub(r'[\\/:*?"<>|\r\n]', "", title).strip()
    return t[:40] or "未命名"


def build_wiki(progress_cb=None, model_choice="auto"):
    """Karpathy 编译：knowledge_docs → wiki/ 结构化词条 + 索引 + 向量检索库
    model_choice: auto(默认快速模型)/reasoning(DeepSeek)/math(API数学模型)"""
    def _report(m):
        if progress_cb is not None:
            try:
                progress_cb(m)
            except Exception:
                pass
        print(m)
    compile_llm, compile_model_name = get_llm_for_choice(model_choice, default="llm_fast")
    _report(f"====LLM Wiki 编译开始：读取 {KNOWLEDGE_FOLDER}，编译模型：{compile_model_name}====")
    support = (".txt", ".md", ".docx", ".pdf")
    files = [f for f in Path(KNOWLEDGE_FOLDER).glob("*")
             if f.is_file() and f.suffix.lower() in support]
    if not files:
        _report("knowledge_docs 无有效文档，请先上传资料")
        return None
    all_items = []
    for fp in files:
        _report(f"读取 {fp.name} ...")
        content = read_knowledge_file(fp)
        if not content.strip():
            continue
        chunks = _split_knowledge_text(content)
        _report(f"  {fp.name} 共 {len(chunks)} 块，逐块编译中")
        for ci, ck in enumerate(chunks):
            _report(f"  ▶ 第 {ci+1}/{len(chunks)} 块提炼词条...")
            arr = None
            for attempt in range(2):  # 解析失败自动重试一次
                try:
                    resp = compile_llm.invoke(
                        [{"role": "user", "content": WIKI_COMPILE_PROMPT.format(chunk=ck)}])
                    m = re.search(r"\[.*\]", resp.content, re.DOTALL)
                    if m:
                        arr = _safe_json_loads(m.group())
                except Exception:
                    arr = None
                if arr:
                    break
                if attempt == 0:
                    _report("    JSON 解析失败，重试一次...")
            if not arr:
                _report("    仍解析失败，跳过该块")
                continue
            for it in arr:
                title = str(it.get("title", "")).strip()
                if not title:
                    continue
                cat = str(it.get("category", "其他")).strip()
                if cat not in CATEGORY_LIST:
                    cat = "其他"
                all_items.append({
                    "title": title,
                    "category": cat,
                    "summary": str(it.get("summary", "")).strip(),
                    "key_points": [str(k).strip() for k in it.get("key_points", []) if str(k).strip()],
                    "related": [str(r).strip() for r in it.get("related", []) if str(r).strip()],
                    "source": fp.name,
                })
    if not all_items:
        _report("未编译出任何词条，请检查资料内容或稍后重试")
        return None
    # —— 写出词条文件（按分类组织）——
    wiki_root = Path(WIKI_FOLDER)
    for sub in wiki_root.iterdir():  # 清空旧编译产物
        if sub.is_dir():
            for f in sub.rglob("*"):
                if f.is_file():
                    f.unlink()
            sub.rmdir()
        elif sub.is_file():
            sub.unlink()
    for cat in CATEGORY_LIST + ["其他"]:
        (wiki_root / cat).mkdir(parents=True, exist_ok=True)
    for it in all_items:
        fn = wiki_root / it["category"] / (_safe_wiki_name(it["title"]) + ".md")
        lines = [f"# {it['title']}", "",
                 f"> 分类：{it['category']} ｜ 来源：{it['source']}", ""]
        if it["summary"]:
            lines += ["## 概述", "", it["summary"], ""]
        if it["key_points"]:
            lines += ["## 要点", ""]
            lines += [f"- {k}" for k in it["key_points"]]
            lines.append("")
        if it["related"]:
            lines += ["## 关联词条", ""]
            lines += [f"- [[{r}]]" for r in it["related"]]
            lines.append("")
        fn.write_text("\n".join(lines), encoding="utf-8")
    # —— 总索引 ——
    index_lines = ["# 📚 LLM Wiki 总索引（由知识库自动编译）", ""]
    index_lines.append(f"> 编译时间：{time.strftime('%Y-%m-%d %H:%M:%S')} ｜ 词条数：{len(all_items)} ｜ 来源文件：{len(files)} 个")
    index_lines.append("")
    for cat in CATEGORY_LIST + ["其他"]:
        items = [it for it in all_items if it["category"] == cat]
        if not items:
            continue
        index_lines.append(f"## {cat}")
        for it in items:
            index_lines.append(f"- [{it['title']}]({cat}/{_safe_wiki_name(it['title'])}.md)")
        index_lines.append("")
    (wiki_root / "index.md").write_text("\n".join(index_lines), encoding="utf-8")
    (wiki_root / "_meta.json").write_text(json.dumps({
        "built_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "count": len(all_items),
        "sources": [f.name for f in files],
    }, ensure_ascii=False, indent=2), encoding="utf-8")
    # —— 建立向量检索库（查询用）——
    _index_wiki_db()
    _report(f"✅ LLM Wiki 编译完成：共 {len(all_items)} 个词条，索引与检索库已更新")
    return all_items


def _index_wiki_db():
    """把 wiki 词条页内容建入 Chroma，供查询检索"""
    wiki_root = Path(WIKI_FOLDER)
    pages = sorted(p for p in wiki_root.rglob("*.md") if p.name != "index.md")
    texts, metas = [], []
    for p in pages:
        texts.append(p.read_text(encoding="utf-8"))
        metas.append(str(p))
    if not texts:
        return None
    embeddings = OllamaEmbeddings(model="nomic-embed-text")
    db = Chroma.from_texts(texts, embedding=embeddings,
                           metadatas=[{"page": m} for m in metas],
                           persist_directory=WIKI_DB_PATH)
    return db


def _query_tokens(query: str):
    """把查询拆成可匹配的词元：英文/数字词 + 中文 2~3 字滑动窗口"""
    toks = set()
    for w in re.findall(r"[a-zA-Z0-9]+", query):
        if len(w) >= 2:
            toks.add(w.lower())
    cn = re.sub(r"[^\u4e00-\u9fa5]", "", query)
    for n in (3, 2):
        for i in range(len(cn) - n + 1):
            toks.add(cn[i:i + n])
    return toks


def wiki_retrieve(query: str, k: int = 3):
    """Karpathy 查询：直接从编译好的 wiki 词条检索。
    关键词打分（标题权重高）优先 + 向量检索补充，符合“直接读 Wiki”思想。
    未编译/失败返回 None（由调用方回退 RAG）。"""
    meta_path = Path(WIKI_FOLDER) / "_meta.json"
    if not meta_path.exists():
        return None
    try:
        pages = [p for p in Path(WIKI_FOLDER).rglob("*.md") if p.name != "index.md"]
        page_data = []
        for p in pages:
            text = p.read_text(encoding="utf-8")
            page_data.append({"title": p.stem, "text": text, "page": str(p)})
        if not page_data:
            return None
        # 1) 关键词打分：命中标题权重 3，命中正文权重 1（n-gram 拆分中文查询）
        keywords = _query_tokens(query)
        scored = []
        for pd in page_data:
            score = 0
            for w in keywords:
                if w in pd["title"]:
                    score += 3
                if w in pd["text"]:
                    score += 1
            scored.append((score, pd))
        scored.sort(key=lambda x: -x[0])
        top = [pd for s, pd in scored if s > 0][:k]
        # 2) 关键词命中不足时，用向量检索补充
        if len(top) < k:
            try:
                embeddings = OllamaEmbeddings(model="nomic-embed-text")
                if os.path.exists(WIKI_DB_PATH) and len(os.listdir(WIKI_DB_PATH)) > 0:
                    db = Chroma(persist_directory=WIKI_DB_PATH, embedding_function=embeddings)
                else:
                    db = _index_wiki_db()
                if db is not None:
                    for d in db.similarity_search(query, k=k):
                        if len(top) >= k:
                            break
                        pg = d.metadata.get("page", "")
                        if pg and not any(pd["page"] == pg for pd in top):
                            top.append({"title": Path(pg).stem,
                                        "text": d.page_content, "page": pg})
            except Exception as e:
                print(f"wiki向量补充检索失败：{e}")
        if not top:
            return None
        context = "\n====[LLM Wiki 知识词条]====\n"
        for pd in top[:k]:
            context += f"\n--- 词条：{pd['title']} ---\n{pd['text']}\n"
        return context
    except Exception as e:
        print(f"wiki检索失败，回退RAG：{e}")
        return None


def wiki_lint(progress_cb=None):
    """Karpathy 自检：检查交叉链接/索引完整性，必要时重建索引"""
    def _report(m):
        if progress_cb is not None:
            try:
                progress_cb(m)
            except Exception:
                pass
        print(m)
    wiki_root = Path(WIKI_FOLDER)
    if not (wiki_root / "_meta.json").exists():
        _report("尚未编译 Wiki，请先点击「编译 Wiki 知识库」")
        return None
    pages = [p for p in wiki_root.rglob("*.md") if p.name != "index.md"]
    index_path = wiki_root / "index.md"
    titles = {p.stem for p in pages}
    issues = []
    link_re = re.compile(r"\[\[([^\]|]+)\]\]")
    # 1) 交叉链接目标
    for p in pages:
        for t in link_re.findall(p.read_text(encoding="utf-8")):
            if t.strip() not in titles:
                issues.append(f"断裂链接：{p.stem} → [[{t}]]")
    # 2) 索引指向
    if index_path.exists():
        for line in index_path.read_text(encoding="utf-8").splitlines():
            m = re.search(r"\]\(([^)]+\.md)\)", line)
            if m and not (wiki_root / m.group(1)).exists():
                issues.append(f"索引失效：{m.group(1)}")
    # 3) 孤儿词条（无入链）
    incoming = {t: 0 for t in titles}
    for p in pages:
        for t in link_re.findall(p.read_text(encoding="utf-8")):
            if t.strip() in incoming:
                incoming[t.strip()] += 1
    orphans = [t for t, c in incoming.items() if c == 0]
    _report(f"词条数：{len(titles)}，交叉链接检查完毕")
    if issues:
        for i in issues:
            _report(f"⚠ {i}")
    else:
        _report("✅ 链接完整，无断裂")
    if orphans:
        _report(f"ℹ 孤儿词条（暂无入链）：{'、'.join(orphans)}")
    if not index_path.exists():
        _rebuild_wiki_index(wiki_root)
        _report("✅ 缺失的 index.md 已重建")
    return issues


def wiki_heal(progress_cb=None):
    """Karpathy 自检修复：扫描所有断裂链接，自动重定向到最匹配词条或降级为纯文本，
    保证 wiki 内不再有指向不存在词条的链接。修复后重建索引。"""
    def _report(m):
        if progress_cb is not None:
            try:
                progress_cb(m)
            except Exception:
                pass
        print(m)
    wiki_root = Path(WIKI_FOLDER)
    if not (wiki_root / "_meta.json").exists():
        _report("尚未编译 Wiki，请先点击「编译 Wiki 知识库」")
        return None
    pages = [p for p in wiki_root.rglob("*.md") if p.name != "index.md"]
    titles = {p.stem for p in pages}
    categories = set(CATEGORY_LIST) | {"其他"}
    link_re = re.compile(r"\[\[([^\]|]+)\]\]")
    redirected, downgraded = 0, 0
    detail = []
    for p in pages:
        text = p.read_text(encoding="utf-8")
        def repl(m):
            nonlocal redirected, downgraded
            target = m.group(1).strip()
            if target in titles:
                return m.group(0)  # 有效链接，保留
            # 1) 目标是分类名（如「导数」「三角」）→ 分类不是词条，降级为纯文本
            if target in categories:
                downgraded += 1
                detail.append(f"降级：{p.stem} → 「{target}」(分类名)")
                return target
            # 2) 目标与某个已有词条标题互为子串 → 重定向到最匹配的词条
            best = None
            for t in titles:
                if target == t:
                    best = t
                    break
                if target in t or t in target:
                    if best is None or len(t) < len(best):
                        best = t
            if best:
                redirected += 1
                detail.append(f"重定向：{p.stem} → 「{target}」→「{best}」")
                return f"[[{best}]]"
            # 3) 无法匹配 → 降级为纯文本
            downgraded += 1
            detail.append(f"降级：{p.stem} → 「{target}」(无匹配)")
            return target
        new_text = link_re.sub(repl, text)
        if new_text != text:
            p.write_text(new_text, encoding="utf-8")
    _rebuild_wiki_index(wiki_root)
    _report(f"✅ Wiki 修复完成：重定向 {redirected} 个链接，降级为纯文本 {downgraded} 个")
    for d in detail:
        _report(f"  · {d}")
    return redirected, downgraded


def _rebuild_wiki_index(wiki_root: Path):
    """按现有词条文件重建总索引"""
    cats = {}
    for p in sorted(wiki_root.rglob("*.md")):
        if p.name == "index.md":
            continue
        cats.setdefault(p.parent.name, []).append(p.stem)
    lines = ["# 📚 LLM Wiki 总索引", ""]
    for cat in CATEGORY_LIST + ["其他"]:
        items = cats.get(cat, [])
        if not items:
            continue
        lines.append(f"## {cat}")
        for t in items:
            lines.append(f"- [{t}]({cat}/{_safe_wiki_name(t)}.md)")
        lines.append("")
    (wiki_root / "index.md").write_text("\n".join(lines), encoding="utf-8")


def wiki_summary() -> str:
    """返回 wiki 编译概况（供网页端展示）"""
    meta_path = Path(WIKI_FOLDER) / "_meta.json"
    if not meta_path.exists():
        return "尚未编译 Wiki（点击「编译 Wiki 知识库」生成）"
    try:
        meta = json.loads(meta_path.read_text(encoding="utf-8"))
        return (f"✅ 已编译 {meta.get('count', 0)} 个词条，"
                f"来源 {len(meta.get('sources', []))} 个文件，"
                f"编译时间 {meta.get('built_at', '')}")
    except Exception:
        return "Wiki 元信息读取失败"

# ========================D 大题自动拆解节点========================
def _looks_like_big_problem(user_q: str) -> bool:
    """粗略判断是否为需要拆解的大题：题内含小问标记（(1)/(1)、1. 等）视为大题"""
    return bool(re.search(r"[（(]\s*\d+\s*[)）]|\d+\s*[.、．]", user_q))

def decompose_problem(state: MessagesState):
    user_q = state["messages"][-1].content
    # 简单题直接跳过拆解（省一次大模型调用，加快速度）
    if not _looks_like_big_problem(user_q):
        return {"messages":[{"role":"system","content":"【题目拆解】无需拆解"}]}
    decompose_prompt = f"""
请把下面这道高考数学大题拆解成若干个子小问题，按顺序输出子问题列表。
如果是简单小题不需要拆解，直接返回【无需拆解】。
题目：{user_q}
输出格式：
是否拆解：是/否
子问题列表：
1. xxx
2. xxx
"""
    resp = llm_fast.invoke([{"role":"user","content":decompose_prompt}])
    return {"messages":[{"role":"system","content":"【题目拆解】"+resp.content}]}

def _recent_context(messages, max_rounds: int = 2):
    """从对话状态中提取最近 max_rounds 轮『用户问+最终答』作为记忆上下文。
    同一问多轮 AI 输出（解题/复核）只保留最后一次；忽略工具调用链与拆解系统消息。"""
    turns = []
    q = None
    for m in messages:
        t = getattr(m, "type", "")
        if t == "human":
            q = m.content
        elif t == "ai" and not getattr(m, "tool_calls", None) and q is not None:
            if turns and turns[-1][0] == q:
                turns[-1] = (q, m.content)  # 同一问的最新回答（复核后版本）
            else:
                turns.append((q, m.content))
    return turns[-max_rounds:]


def _current_question_and_decompose(state):
    """返回 (当前用户题目, 当前题目的拆解note)"""
    msgs = state["messages"]
    user_q = None
    for m in reversed(msgs):
        if getattr(m, "type", "") == "human":
            user_q = m.content
            break
    decompose_note = None
    for m in reversed(msgs):
        if getattr(m, "type", "") == "system" and str(m.content).startswith("【题目拆解】"):
            decompose_note = m.content
            break
    return user_q, decompose_note


# 指代词：命中即认为当前问题可能引用上一题，需要把上一题拼进来保证题目自包含
_REF_WORDS = re.compile(r"它|该|上题|上述|此|这个|结果|继续|其|之|本题|前面|刚才")


def _resolve_user_question(state):
    """若当前问题含指代词且存在上一题，则把上一题题目拼进当前问题，使其自包含。"""
    user_q, decompose_note = _current_question_and_decompose(state)
    resolved = user_q or ""
    if resolved and _REF_WORDS.search(resolved):
        recent = _recent_context(state["messages"], max_rounds=1)
        if recent:
            prev_q = recent[-1][0]
            resolved = f"[本题承接上一题：{prev_q}]\n当前问题：{resolved}"
    return resolved, decompose_note


def solve_agent(state: MessagesState):
    user_content, decompose_note = _resolve_user_question(state)
    rag_ctx = rag_retrieve(user_content)
    system_prompt = f"""
你是高考数学解题专家，请输出一份「完整、规范、精炼」的高考标准答案（类似官方评分标准答案）。

参考知识库信息（仅作知识点参考）：
{rag_ctx}
硬性规则：
1. 【强制计算器】所有涉及数值计算的步骤（求导、解方程、不等式、向量点乘/叉乘/模长、距离、角度、面积、体积、概率等）必须先调用 sympy_math_calc 工具得到准确结果，禁止心算给出数值结果！调用工具时写清楚计算表达式（如向量叉乘：sympy.Matrix([1,2,3]).cross(sympy.Matrix([4,5,6]))；距离：sympy.sqrt((x2-x1)**2+(y2-y1)**2+(z2-z1)**2)）。若工具调用失败，重试一次；仍失败才给出你计算的结果并标注"（心算，建议验证）"。
2. 用户需要图像时调用 plot_function 绘图。
3. 禁止把 python 代码写在回答文本中，代码只能出现在工具调用里。
4. 【重要】禁止输出 Markdown 和 LaTeX 语法（不要出现 \\[、\\]、\\(、\\)、\\frac、\\dfrac、\\boxed、\\sqrt、**、#、\\ln 等）。
   数学式一律用普通文本 + Unicode 表示，例如：
   - 导数写为：f'(x) = ln x + 2 - 2a x
   - 分数写为：(a)/(b) 或 a/b
   - 幂/下标写为：x^2、x1、x2、e/2
   - 取值范围写为：0 < a < e/2
5. 输出正文就是标准答案本身，不要出现“解题思路”“考点总结”“教学点评”等多余板块，不要自我重复。
6. 【确定正确】禁止出现“似乎”“可能”“也许”“需要进一步验证”“这没有直接帮助”等探索性或不确定表述；每一步都要给出确定、正确的结论，若有疑问就独立重新推导出确定结果。
7. 【详略得当】步骤要完整、详细：不跳步，写清楚每一步的依据（如“由求导法则”“由对数性质”“令导数等于零”等）和中间结果；简单题也必须写完整推导，禁止只写结论；但不要写与解题无关的解释性文字。
8. 高频易错提醒：过点求切线先验证点是否在曲线上（不在则设切点法）；不等式乘除含字母式子先判正负（负则变号）；圆锥曲线联立先写判别式再用韦达定理。
9. 【立体几何专项】立体几何题必须建立空间直角坐标系，写出各点坐标后用向量法求解：
   - 证明平行/垂直：用向量点乘为0（垂直）或叉乘为0（平行），必须调用计算器验证点乘/叉乘结果
   - 求线面角：先求平面法向量（用叉乘，必须调用计算器），线面角=arcsin(|方向向量·法向量|/(|方向向量|*|法向量|))
   - 求二面角：求两个面的法向量（叉乘），二面角余弦=法向量点乘/(|法向量1|*|法向量2|)，注意判断锐角还是钝角
   - 求点到平面距离：|点·法向量 - d| / |法向量|，必须调用计算器计算
   - 求体积：底面积×高/3，各量必须调用计算器验证
   所有向量运算（点乘、叉乘、模长）和最终数值结果必须调用 sympy_math_calc 工具，禁止心算！

输出结构（按题目小问顺序）：
每小问：【解】以编号步骤逐条完整推导——写成“步骤1：… 步骤2：…”，每一步写明依据、公式变换和中间结果，逐步推进到结论；步骤详略适中（一般 3～8 步）。
最后：【答案】汇总各小问的最终结论，简洁明确。
参考拆解出的子问题，分小问依次完整作答。
"""
    # 只把『当前题（含解析后的指代上下文）+ 拆解note』喂给模型，不携带全部历史
    messages = [{"role": "system", "content": system_prompt}]
    if decompose_note:
        messages.append({"role": "system", "content": decompose_note})
    messages.append({"role": "user", "content": user_content})
    # 按题型自动路由 / 手动选择模型
    choice = state.get("model_choice", "auto")
    model_tools, model_plain, model_name = resolve_model(choice, user_content)
    ai_msg = model_tools.invoke(messages)
    return {"messages": [ai_msg], "model_used": model_name}

def checker_agent(state: MessagesState):
    # 使用解析后的完整题目（含承接的上一题），确保复核时不会丢失指代上下文
    user_content, _ = _resolve_user_question(state)
    last_solve = ""
    for m in state["messages"]:
        if getattr(m, "type", "") == "ai" and not getattr(m, "tool_calls", None):
            last_solve = m.content
    check_prompt = f"""
你是数学阅卷老师。下面是题目和初步解答，请独立重新演算，检查计算错误和逻辑漏洞。

【题目】
{user_content}

【初步解答】
{last_solve}

规则：
- 只输出最终的标准答案正文（复核修正后的完整标准答案），你是最终标准答案的权威输出者。
- 若初步解答有误，输出修正后完整、正确的标准答案；若正确，直接输出同一份标准答案。
- 禁止输出任何“复核”“校验”“检查过程”“评语”，禁止重复上面的题目和初步解答原文。
- 禁止 Markdown 和 LaTeX 语法（不要出现 \\[、\\]、\\(、\\)、\\frac、\\boxed、**、# 等），数学式一律用普通文本 + Unicode 表示。
- 按题目小问顺序给出【解】和完整的逐步推导：用“步骤1：… 步骤2：…”逐条编号，每一步写明依据、公式变换和中间结果，不跳步；简单题也必须写完整推导，禁止只写结论。
- 每条步骤必须确定、正确：禁止出现“似乎”“可能”“需要进一步验证”等探索性或不确定表述；若有疑问，独立重新计算得出确定结论，不要写出探索过程。
- 【验算要求】对解答中的所有关键数值计算（向量点乘/叉乘/模长、距离、角度、面积、体积、导数、方程的根、概率等）必须独立重新演算一遍，若发现原解答计算错误，输出修正后的正确答案；立体几何题必须重新验证坐标系建立、各点坐标、法向量计算和最终数值结果。
- 步骤详略适中（一般 3～8 步），不写与解题无关的文字。最后给出【答案】汇总各小问结论。
"""
    choice = state.get("model_choice", "auto")
    _, model_plain, model_name = resolve_model(choice, user_content)
    resp = model_plain.invoke([{"role": "user", "content": check_prompt}])
    return {"messages": [resp], "model_used": model_name}

def clean_model_answer(text: str) -> str:
    """清理模型输出中的 DeepSeek 特殊标记与 Markdown/LaTeX 残留，避免 Word 文档出现乱码和重复"""    # 1) 去掉 deepseek-r1 的 response / thinking 特殊标记；
    #    若内容中存在多个 response 标记（模型复读导致答案重复），只保留最后一段
    parts = re.split(r'\s*<\|?/?response\|?>?\s*', text, flags=re.IGNORECASE)
    text = parts[-1] if len(parts) > 1 else text
    text = re.sub(r'^\s*<\|?/?think(ing)?\|?>?\s*', '', text, flags=re.IGNORECASE)
    # 2) 去掉 Markdown 标题符号（行首 # / ## 等）和强调符号 **
    text = re.sub(r'(?m)^\s*#{1,6}\s*', '', text)
    text = text.replace('**', '')
    # 3) 去掉 LaTeX 公式定界符残留
    text = text.replace('\\[', '').replace('\\]', '')
    text = text.replace('\\(', '').replace('\\)', '')
    # 4) LaTeX 数学符号 → 纯文本（供 Word 正常显示）
    text = latex_to_plain(text)
    # 5) 清理每行首尾空白
    lines = [ln.rstrip() for ln in text.split('\n')]
    return '\n'.join(lines).strip()


def run_question(graph, question: str, config, max_tries: int = 2, model_choice: str = "auto"):
    """统一答题入口：调用图解题并清洗输出；若模型只输出推理、正文为空，自动重试。
    model_choice: "auto"(按题型自动路由) / "reasoning"(DeepSeek) / "math"(Qwen数学)。
    返回 (清洗后答案, 实际使用的模型名)。"""
    ans = ""
    model_used = "未知"
    for attempt in range(max_tries):
        res = graph.invoke(
            {"messages": [{"role": "user", "content": question}], "model_choice": model_choice},
            config=config)
        ans = clean_model_answer(res["messages"][-1].content)
        model_used = res.get("model_used", "未知")
        if len(ans) >= 20:
            if "siliconflow" in model_used.lower():
                record_api_usage()
            return ans, model_used
        print(f"⚠️ 本次未生成有效答案（模型可能只输出了推理过程），第{attempt + 1}次重试中...")
    return ans, model_used


def _find_matching_brace(s: str, start: int) -> int:
    """s[start] == '{'，返回与之匹配的 '}' 下标（用于 LaTeX 花括号配对）"""
    depth = 0
    for i in range(start, len(s)):
        if s[i] == '{':
            depth += 1
        elif s[i] == '}':
            depth -= 1
            if depth == 0:
                return i
    return len(s) - 1


# LaTeX 简单命令 → Unicode/纯文本 映射
_LATEX_SYMBOL_MAP = {
    'ln': 'ln', 'log': 'log', 'sin': 'sin', 'cos': 'cos', 'tan': 'tan',
    'ge': '≥', 'le': '≤', 'neq': '≠', 'to': '→', 'rightarrow': '→',
    'in': '∈', 'pi': 'π', 'pm': '±', 'times': '×', 'div': '÷',
    'mid': '|', 'cdot': '*', 'implies': '=>', 'infty': '∞',
    'text': '', 'mathrm': '', 'operatorname': '', 'quad': '  ', 'qquad': '  ',
    'left': '', 'right': '', 'big': '', 'Big': '', 'bigg': '', 'Bigg': '',
}


def latex_to_plain(s: str) -> str:
    """把常用 LaTeX 数学式转成可读纯文本（供 Word 无乱码显示）"""
    out = []
    i = 0
    n = len(s)
    while i < n:
        c = s[i]
        if c == '\\':
            j = i + 1
            while j < n and s[j].isalpha():
                j += 1
            cmd = s[i + 1:j]
            low = cmd.lower()
            # \frac{a}{b} / \dfrac{a}{b} / \tfrac{a}{b}
            if low in ('frac', 'dfrac', 'tfrac', 'cfrac'):
                num = den = ''
                k = j
                while k < n and s[k] == ' ': k += 1
                if k < n and s[k] == '{':
                    e = _find_matching_brace(s, k)
                    num = latex_to_plain(s[k + 1:e])
                    k = e + 1
                while k < n and s[k] == ' ': k += 1
                if k < n and s[k] == '{':
                    e = _find_matching_brace(s, k)
                    den = latex_to_plain(s[k + 1:e])
                    k = e + 1
                out.append(f'({num})/({den})')
                i = k
                continue
            # \sqrt{...}
            if low == 'sqrt':
                k = j
                while k < n and s[k] == ' ': k += 1
                if k < n and s[k] == '{':
                    e = _find_matching_brace(s, k)
                    out.append('sqrt(' + latex_to_plain(s[k + 1:e]) + ')')
                    i = e + 1
                else:
                    out.append('sqrt')
                    i = k
                continue
            # \boxed{...} 保留内部内容
            if low == 'boxed':
                k = j
                while k < n and s[k] == ' ': k += 1
                if k < n and s[k] == '{':
                    e = _find_matching_brace(s, k)
                    out.append(latex_to_plain(s[k + 1:e]))
                    i = e + 1
                    continue
            # 其它已知命令
            if low in _LATEX_SYMBOL_MAP:
                out.append(_LATEX_SYMBOL_MAP[low])
                i = j
                continue
            # 未知命令：丢弃
            i = j
            continue
        elif c == '{' or c == '}':
            i += 1
            continue
        elif c == '^':
            # 上标：x^{2} -> x^2；e^{-1} -> e^(-1)
            k = i + 1
            while k < n and s[k] == ' ': k += 1
            if k < n and s[k] == '{':
                e = _find_matching_brace(s, k)
                sup = latex_to_plain(s[k + 1:e])
                i = e + 1
            else:
                sup = s[k] if k < n else ''
                i = k + 1
            if any(ch in sup for ch in '+-*/()'):
                out.append('^(' + sup + ')')
            else:
                out.append('^' + sup)
            continue
        elif c == '_':
            # 下标：x_1 -> x1
            k = i + 1
            while k < n and s[k] == ' ': k += 1
            if k < n and s[k] == '{':
                e = _find_matching_brace(s, k)
                out.append(latex_to_plain(s[k + 1:e]))
                i = e + 1
            elif k < n:
                out.append(s[k])
                i = k + 1
            else:
                i += 1
            continue
        else:
            out.append(c)
            i += 1
    return ''.join(out)

tool_node = ToolNode(tools)

# 复核开关：默认开启（保证标准答案质量）；可在 .env 设 ENABLE_CHECKER=off 关闭以提速
_ENABLE_CHECKER = os.getenv("ENABLE_CHECKER", "on").lower() in ("1", "true", "on", "yes")

def should_continue(state: MessagesState):
    last_msg = state["messages"][-1]
    if last_msg.tool_calls:
        return "tools"
    if _ENABLE_CHECKER:
        return "checker"
    return "__end__"

builder = StateGraph(SolveState)
builder.add_node("decompose", decompose_problem)
builder.add_node("solve_agent", solve_agent)
builder.add_node("tools", tool_node)
builder.add_node("checker_agent", checker_agent)
builder.set_entry_point("decompose")
builder.add_edge("decompose", "solve_agent")
builder.add_conditional_edges(
    "solve_agent", should_continue,
    {"tools": "tools", "checker": "checker_agent", "__end__": "__end__"})
builder.add_edge("tools", "solve_agent")
builder.add_edge("checker_agent", "__end__")

# langgraph-checkpoint-sqlite 3.x 中 from_conn_string 已改为上下文管理器，
# 这里直接构造 SqliteSaver(conn)，连接在进程生命周期内保持，实现跨会话记忆持久化。
conn = sqlite3.connect("./all_memory.db", check_same_thread=False)
memory = SqliteSaver(conn)
graph_chat = builder.compile(checkpointer=memory)

# ======================== 批量做题模块｜支持PDF输入，输出docx，不再输出md ========================
def _repair_extracted_text(text: str) -> str:
    """修复PDF/文档提取中的常见文本问题：行断裂、题号分离、选项混入等。"""
    if not text or len(text.strip()) < 5:
        return text
    lines = text.split("\n")
    repaired = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        # 情况1：行尾是题号（如"1."），下一行是题目内容 → 合并
        if re.match(r"^\d+\s*[.、．]\s*$", line) and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if next_line and not re.match(r"^\d+\s*[.、．]", next_line):
                repaired.append(line + " " + next_line)
                i += 2
                continue
        # 情况2：行尾不是结束标点，且下一行不是新题/选项开头 → 合并（修复断行）
        if i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            ends_with_terminal = re.search(r"[。！？；：）\]】]$", line)
            next_is_new_question = re.match(r"^\d+\s*[.、．]", next_line)
            next_is_option = re.match(r"^[A-D]\s*[.、．]", next_line)
            next_is_section = re.match(r"^[一二三四五六七八九十]+、", next_line)
            if (not ends_with_terminal and next_line
                    and not next_is_new_question and not next_is_option
                    and not next_is_section and len(next_line) > 0):
                repaired.append(line + next_line)
                i += 2
                continue
        repaired.append(line)
        i += 1
    result = "\n".join(repaired)
    # 清理多余空白
    result = re.sub(r"[ \t]+", " ", result)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip()


def _pdf_ocr_fallback(doc) -> str:
    """OCR兜底：对扫描版PDF逐页OCR提取文本。"""
    text_all = ""
    try:
        for page in doc:
            tp = page.get_textpage_ocr(flags=3, language="chi_sim+eng", dpi=300, full=True)
            pt = page.get_text(textpage=tp)
            if pt:
                text_all += pt + "\n"
    except Exception as e:
        print(f"OCR失败: {e}")
    return text_all


def read_batch_file(filepath:Path):
    """读取batch_input支持：txt md docx pdf"""
    suffix = filepath.suffix.lower()
    if suffix == ".pdf":
        text_all = ""
        doc = fitz.open(filepath)
        for page in doc:
            pt = page.get_text("text")
            if pt:
                text_all += pt + "\n"
        doc.close()
        # 文本修复
        text_all = _repair_extracted_text(text_all)
        # OCR兜底：提取文本过少（可能是扫描版）
        if len(text_all.strip()) < 100:
            print("PDF文本提取量过少，尝试OCR...")
            doc = fitz.open(filepath)
            ocr_text = _pdf_ocr_fallback(doc)
            doc.close()
            if len(ocr_text.strip()) > len(text_all.strip()):
                text_all = _repair_extracted_text(ocr_text)
                print(f"OCR提取成功，共{len(text_all)}字符")
        return text_all
    elif suffix == ".docx":
        doc = Document(filepath)
        raw = "\n".join([p.text for p in doc.paragraphs])
        return _repair_extracted_text(raw)
    elif suffix in (".txt",".md"):
        with open(filepath,"r",encoding="utf-8") as f:
            raw = f.read()
        return _repair_extracted_text(raw)
    else:
        raise Exception(f"不支持文件后缀 {suffix}")

def _docx_to_md(filepath: Path) -> str:
    """将DOCX转换为Markdown（基于python-docx，保留标题/段落/表格结构）。"""
    from docx import Document
    doc = Document(filepath)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        style = para.style.name if para.style else ""
        if style.startswith("Heading 1") or "标题 1" in style or "标题1" in style:
            lines.append(f"# {text}")
        elif style.startswith("Heading 2") or "标题 2" in style or "标题2" in style:
            lines.append(f"## {text}")
        elif style.startswith("Heading 3") or "标题 3" in style or "标题3" in style:
            lines.append(f"### {text}")
        elif "List" in style or "列表" in style:
            lines.append(f"- {text}")
        else:
            lines.append(text)
    # 追加表格（简单转换）
    for table in doc.tables:
        if not table.rows:
            continue
        lines.append("")
        header = table.rows[0]
        lines.append("| " + " | ".join(c.text.strip().replace("\n", " ") for c in header.cells) + " |")
        lines.append("| " + " | ".join("---" for _ in header.cells) + " |")
        for row in table.rows[1:]:
            lines.append("| " + " | ".join(c.text.strip().replace("\n", " ") for c in row.cells) + " |")
        lines.append("")
    return "\n".join(lines).strip()


def convert_file_to_md(filepath, output_folder=None):
    """将文件转换为Markdown，返回(md文本, md文件路径)。
    支持 pdf / docx / txt / md。"""
    filepath = Path(filepath)
    if output_folder is None:
        output_folder = Path("output_md")
    output_folder = Path(output_folder)
    output_folder.mkdir(exist_ok=True)
    suffix = filepath.suffix.lower()
    if suffix == ".pdf":
        import pymupdf4llm
        md_text = pymupdf4llm.to_markdown(str(filepath))
    elif suffix == ".docx":
        md_text = _docx_to_md(filepath)
    elif suffix in (".txt", ".md"):
        with open(filepath, "r", encoding="utf-8") as f:
            md_text = f.read()
    else:
        raise Exception(f"不支持的格式: {suffix}，支持 pdf/docx/txt/md")
    md_path = output_folder / (filepath.stem + ".md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_text)
    return md_text, str(md_path)


def _clean_question_text(text: str) -> str:
    """清理题目文本尾部附带的大题说明、答题要求等非题目内容。"""
    # 截掉从"二、选择题"/"三、填空题"等大题说明开始的尾部内容
    text = re.sub(r"\n?[一二三四五六七八九十]+、\s*(选择题|多选题|填空题|解答题)[\s\S]*$", "", text)
    # 截掉答题要求说明
    text = re.sub(r"在每小题给出的四个选项中[\s\S]*$", "", text)
    text = re.sub(r"只有一项是符合题目要求[\s\S]*$", "", text)
    text = re.sub(r"有多项符合题目要求[\s\S]*$", "", text)
    return text.strip()


def _is_exam_question(text: str) -> bool:
    """判断一段文本是否是真正的数学题目，过滤试卷标题、题型说明等。"""
    text_stripped = text.strip()
    if len(text_stripped) < 10:
        return False
    # 只过滤明显的试卷标题/页眉/大题题型说明
    header_patterns = [
        r"普通高等学校招生", r"全国统一考试", r"高考试题", r"高考真题",
        r"^[一二三四五六七八九十]+、\s*(选择题|填空题|解答题|多选题)",
        r"^第[\u2160-\u2163I]+卷", r"^注意事项", r"^考生须知",
        r"在每小题给出的四个选项中", r"只有一项是符合题目要求",
    ]
    for pat in header_patterns:
        if re.search(pat, text_stripped):
            return False
    return True


def batch_solve_file(filename: str, progress_cb=None, model_choice="auto", resume=True):
    """批量解答一份试卷。支持断点续跑：每解完一题自动保存进度，中断后重新运行会跳过已完成的题。
    model_choice: auto/reasoning/math；resume: True=读取进度续跑，False=从头开始。"""
    def _report(msg):
        if progress_cb is not None:
            try:
                progress_cb(msg)
            except Exception:
                pass
        print(msg)
    file_path = Path(BATCH_INPUT_FOLDER)/filename
    content = read_batch_file(file_path)
    # 按“行首编号”切分多道习题：支持 1. / 1、/ 1． 三种编号；
    # 只匹配行首的数字编号，避免把题内小问 (1)(2) 或文本中的小数误拆。
    pattern = re.compile(r"(?m)^\s*\d+\s*[.、．]\s*")
    parts = pattern.split(content)
    raw_list = [p.strip() for p in parts if len(p.strip())>3]
    # 先清理每题尾部附带的大题说明，再过滤非题目内容
    q_list = []
    for item in raw_list:
        cleaned = _clean_question_text(item)
        if _is_exam_question(cleaned):
            q_list.append(cleaned)
        else:
            print("跳过非题目内容：" + cleaned[:50].replace("\n", " "))
    progress_file = Path(OUTPUT_FOLDER) / (Path(filename).stem + "_progress.json")
    all_out_items = []
    start_idx = 0
    if resume and progress_file.exists():
        try:
            with open(progress_file, "r", encoding="utf-8") as f:
                saved = json.load(f)
            all_out_items = saved.get("items", [])
            start_idx = len(all_out_items)
            if 0 < start_idx < len(q_list):
                _report("检测到进度，已完成 " + str(start_idx) + "/" + str(len(q_list)) + " 题，从第 " + str(start_idx+1) + " 题续跑")
        except Exception:
            all_out_items = []
            start_idx = 0
    thread_id = "batch_session"
    config = {"configurable":{"thread_id":thread_id}, "recursion_limit": 10}
    graph_chat.update_state(config, {"messages":[]})
    if start_idx == 0:
        _report("共识别 " + str(len(q_list)) + " 道题，从头开始")
    for idx in range(start_idx, len(q_list)):
        q = q_list[idx]
        _report("正在解答第 " + str(idx+1) + "/" + str(len(q_list)) + " 题：" + q[:30] + "...")
        print("\n====批量处理第" + str(idx+1) + "题====\n题目：" + q[:80])
        try:
            ans_text, _model_used = run_question(graph_chat, q, config, model_choice=model_choice)
        except Exception as e:
            ans_text = "[本题解答出错：" + type(e).__name__ + "]"
            _report("第" + str(idx+1) + "题出错，已记录并继续")
        is_error = ans_text.startswith("[本题解答出错")
        all_out_items.append({
            "no": idx+1,
            "question": q,
            "answer": ans_text,
            "success": not is_error
        })
        if is_error:
            _report("第" + str(idx+1) + "题解答出错，已标记，继续下一题")
        try:
            with open(progress_file, "w", encoding="utf-8") as f:
                json.dump({"filename": filename, "total": len(q_list), "items": all_out_items}, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    # 统计成功/失败
    success_count = sum(1 for it in all_out_items if it.get("success", True))
    fail_count = len(all_out_items) - success_count
    fail_nos = [str(it["no"]) for it in all_out_items if not it.get("success", True)]
    _report("第一遍完成：共" + str(len(all_out_items)) + "题，成功" + str(success_count) + "题，失败" + str(fail_count) + "题")

    # —— 自动重做错题（第二遍）——
    retry_fixed = []
    if fail_count > 0:
        _report("检测到" + str(fail_count) + "道失败题，开始自动重做（第二遍，每题最多重试3次）...")
        graph_chat.update_state(config, {"messages": []})
        for it in all_out_items:
            if not it.get("success", True):
                qno = it["no"]
                q = it["question"]
                _report("重做第" + str(qno) + "题：" + q[:30] + "...")
                print("\n====重做第" + str(qno) + "题====\n题目：" + q[:80])
                try:
                    retry_ans, _retry_model = run_question(graph_chat, q, config, max_tries=3, model_choice=model_choice)
                    if len(retry_ans) >= 20 and not retry_ans.startswith("[本题解答出错"):
                        it["answer"] = retry_ans
                        it["success"] = True
                        it["retry_fixed"] = True
                        retry_fixed.append(str(qno))
                        _report("第" + str(qno) + "题重做成功！")
                    else:
                        _report("第" + str(qno) + "题重做仍失败，保留原错误标记")
                except Exception as e:
                    _report("第" + str(qno) + "题重做出错：" + str(e)[:100])
                # 保存进度
                try:
                    with open(progress_file, "w", encoding="utf-8") as f:
                        json.dump({"filename": filename, "total": len(q_list), "items": all_out_items}, f, ensure_ascii=False, indent=2)
                except Exception:
                    pass

    # 重新统计（重做后）
    success_count = sum(1 for it in all_out_items if it.get("success", True))
    fail_count = len(all_out_items) - success_count
    fail_nos = [str(it["no"]) for it in all_out_items if not it.get("success", True)]
    if retry_fixed:
        _report("第二遍重做修复了" + str(len(retry_fixed)) + "题：" + "、".join(retry_fixed))
    _report("最终统计：共" + str(len(all_out_items)) + "题，成功" + str(success_count) + "题，失败" + str(fail_count) + "题")
    if fail_nos:
        _report("仍失败的题：" + "、".join(fail_nos) + "（已在文档末尾标注，建议手动检查）")
    else:
        _report("全部题目解答成功！")

    # 直接生成docx，不再生成md
    doc = Document()
    doc.add_heading(f"试卷解答：{filename}", level=1)
    for item in all_out_items:
        status_tag = ""
        if not item.get("success", True):
            status_tag = " ⚠️解答出错"
        elif item.get("retry_fixed", False):
            status_tag = " ✅重做修复"
        doc.add_heading(f"第{item['no']}题{status_tag}", level=2)
        doc.add_paragraph(f"【题目】{item['question']}")
        doc.add_paragraph("【解答】")
        for line in item["answer"].split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.add_paragraph("-"*60)

    # 文档末尾加错题汇总
    doc.add_page_break()
    doc.add_heading("答题情况汇总", level=1)
    doc.add_paragraph("总题数：" + str(len(all_out_items)))
    doc.add_paragraph("成功：" + str(success_count) + "题")
    doc.add_paragraph("失败：" + str(fail_count) + "题")
    if retry_fixed:
        doc.add_paragraph("第二遍自动重做修复：" + str(len(retry_fixed)) + "题（第" + "、".join(retry_fixed) + "题）")
    if fail_nos:
        doc.add_paragraph("")
        doc.add_heading("失败题目汇总（建议单独重做）", level=2)
        for it in all_out_items:
            if not it.get("success", True):
                doc.add_paragraph("第" + str(it["no"]) + "题：" + it["question"][:80] + "...")
                doc.add_paragraph("错误信息：" + it["answer"])
                doc.add_paragraph("-" * 40)
    else:
        doc.add_paragraph("全部题目解答成功！")

    docx_save = Path(OUTPUT_FOLDER) / f"{Path(filename).stem}_result.docx"
    doc.save(str(docx_save))
    print("批量解答保存docx：" + str(docx_save.resolve()))
    try:
        if progress_file.exists():
            progress_file.unlink()
    except Exception:
        pass

    # 返回纯文本，用于后续生成pdf
    full_text_buf = []
    for it in all_out_items:
        full_text_buf.append(f"## 第{it['no']}题\n【题目】{it['question']}\n【解答】\n{it['answer']}\n\n")
    full_text = "\n".join(full_text_buf)
    return full_text, str(docx_save)

def text_to_pdf(text:str, pdf_save_path:str):
    c = canvas.Canvas(pdf_save_path, pagesize=A4)
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    c.setFont("STSong-Light",11)
    width, height = A4
    y_pos = height-40
    lines = text.split("\n")
    for line in lines:
        if y_pos < 40:
            c.showPage()
            c.setFont("STSong-Light",11)
            y_pos = height-40
        c.drawString(30, y_pos, line[:120])
        y_pos -=18
    c.save()
    print(f"✅PDF输出：{pdf_save_path}")

# ========================教研模块：PDF/Word真题自动教研========================
# MathType/Symbol 字体的私有区乱码符号 → 正确 Unicode 数学符号
_SYMBOL_MAP = {
    '\uf028':'(', '\uf029':')', '\uf02b':'+', '\uf02d':'-', '\uf03c':'<', '\uf03d':'=', '\uf03e':'>',
    '\uf05e':'\u22a5',  # ⊥
    '\uf061':'\u03b1', '\uf062':'\u03b2', '\uf070':'\u03c0', '\uf073':'\u03c3',
    '\uf07b':'{', '\uf07c':'|', '\uf07d':'}',
    '\uf0a2':'\u2032', '\uf0a3':'\u2264', '\uf0a5':'\u221e', '\uf0b3':'\u2265', '\uf0bb':'\u2248',
    '\uf0e6':'(', '\uf0e7':'(', '\uf0e8':')',
    '\uf0ec':'{', '\uf0ed':'{', '\uf0ee':'}',
    '\uf0f6':')', '\uf0f7':')', '\uf0f8':')',
    '\uf049':'\u2229',  # ∩
    '\uf072':'\u2192',  # →
    '\uf056':'\u25b3',  # △
}

def clean_math_symbols(text: str) -> str:
    """把 MathType 公式转成的私有区乱码还原为正常数学符号。"""
    return ''.join(_SYMBOL_MAP.get(ch, ch) for ch in text)

def read_pdf(filepath: Path) -> str:
    text_all = ""
    doc = fitz.open(filepath)
    for page in doc:
        page_text = page.get_text("text")
        if page_text:
            text_all += clean_math_symbols(page_text) + "\n"
    doc.close()
    return text_all

def read_docx(filepath: Path) -> str:
    doc = Document(filepath)
    paras = [p.text for p in doc.paragraphs]
    return "\n".join(paras)

def load_raw_exam_files() -> list[tuple[str,str]]:
    res = []
    for f in Path(RAW_EXAM_FOLDER).glob("*"):
        suffix = f.suffix.lower()
        if suffix == ".pdf":
            content = read_pdf(f)
            res.append((f.name, content))
        elif suffix == ".docx":
            content = read_docx(f)
            res.append((f.name, content))
    return res

EXTRACT_PROMPT = """
下面是高考数学试卷文本，请把里面所有数学题目完整提取出来。
注意：同一大题的多个小问（如 (1)(2)(3)）必须合并为同一条记录，题目的 question 字段要完整包含该大题题干及其全部小问。
输出格式JSON数组，每一条：
{{
  "question": "完整原题文本（含该大题全部小问）",
  "category": "从【集合、函数、导数、三角、数列、立体几何、圆锥曲线、概率统计】选择最合适的一个分类"
}}
只输出JSON，不要多余解释。
试卷文本：
{text_chunk}
"""

RESEARCH_PROMPT = """
你是高中数学教研专家。给定一道高考真题，请完整输出下面全部模块，不要省略。
【原题】
{q_text}

输出严格包含下面全部板块：
1.【标准答案】
2.【详细解答】分步完整推导
3.【解题方法】本题用到的核心解题思路与方法
4.【第二种解法】提供另外一套不同思路的完整解法，如果没有写“本题无第二种可行解法”
5.【易错点】学生高频踩坑、容易错在哪里
6.【命题特点】高考命题角度、考察知识点、能力要求
7.【一题三变】给出3道变式题，改变条件，同源考点，用于训练。
"""

def _safe_json_loads(s: str):
    """宽容解析 AI 输出的 JSON 数组：自动清理非法控制字符，避免整段题目被丢弃。"""
    # 尝试直接解析
    try:
        return json.loads(s)
    except json.JSONDecodeError:
        pass
    # 非法控制字符（JSON 任何位置都不允许 0x00-0x08/0x0b/0x0c/0x0e-0x1f/0x7f）直接移除
    s2 = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', s)
    try:
        return json.loads(s2)
    except json.JSONDecodeError:
        pass
    # 字符串内部未转义的换行/制表符替换为空格（牺牲字符串内换行，换取整段可解析）
    s3 = s2.replace('\t', ' ').replace('\r', ' ')
    s3 = re.sub(r'(?<=[^"\\])\n(?=[^"])', ' ', s3)
    try:
        return json.loads(s3)
    except json.JSONDecodeError:
        return None


def research_workflow(only_files=None, progress_cb=None, model_choice="auto"):
    """教研汇总：从 raw_exam_files 提取题目并做教研分析。
    model_choice: auto/reasoning/math，控制题目提取和教研分析使用的模型"""
    def _report(msg):
        if progress_cb is not None:
            try:
                progress_cb(msg)
            except Exception:
                pass
        print(msg)
    research_llm, research_model_name = get_llm_for_choice(model_choice, default="llm")
    _report(f"====教研模块，读取文件夹 {RAW_EXAM_FOLDER}，分析模型：{research_model_name}====")
    file_list = load_raw_exam_files()
    if only_files:
        only_set = set(only_files)
        file_list = [x for x in file_list if x[0] in only_set]
    if len(file_list)==0:
        _report("请把PDF/docx真题放入raw_exam_files")
        return
    _report(f"读取到 {len(file_list)} 份试卷")
    all_research_items = []
    for fname,full_text in file_list:
        _report(f"\n【{fname}】开始处理")
        chunks = re.split(r"\n{3,}",full_text)
        valid_chunks = [c.strip() for c in chunks if len(c.strip())>=80]
        _report(f"【{fname}】共 {len(valid_chunks)} 个有效片段")
        for ci, ck in enumerate(valid_chunks):
            _report(f"【{fname}】正在分析第 {ci+1}/{len(valid_chunks)} 段...")
            try:
                resp_ext = research_llm.invoke([{"role":"user","content":EXTRACT_PROMPT.format(text_chunk=ck)}])
                raw = resp_ext.content
            except Exception as _ext_err:
                _report(f"【{fname}】第 {ci+1}/{len(valid_chunks)} 段提取调用失败（{type(_ext_err).__name__}），跳过本段继续")
                continue
            json_match = re.search(r"\[.*\]", raw, re.DOTALL)
            if not json_match:
                _report(f"【{fname}】第 {ci+1}/{len(valid_chunks)} 段未提取到题目，跳过")
                continue
            try:
                q_arr = _safe_json_loads(json_match.group())
                if q_arr is None:
                    _report(f"【{fname}】第 {ci+1}/{len(valid_chunks)} 段解析失败（JSON无法修复），跳过")
                    continue
                for item in q_arr:
                    q = item.get("question","")
                    cat = item.get("category","未知分类")
                    _report(f"  >>教研题目 {q[:40]}... 分类:{cat}")
                    ana_resp = research_llm.invoke([{"role":"user","content":RESEARCH_PROMPT.format(q_text=q)}])
                    all_research_items.append({
                        "category":cat,
                        "question":q,
                        "research_content":clean_model_answer(ana_resp.content)
                    })
            except Exception as e:
                _report(f"【{fname}】片段解析跳过 {e}")
            _report(f"【{fname}】第 {ci+1}/{len(valid_chunks)} 段完成")
        _report(f"【{fname}】处理完成")
    #导出word
    doc = Document()
    doc.add_heading("高考数学真题教研资料", level=1)
    group = {c:[] for c in CATEGORY_LIST}
    group["其他"]=[]
    for it in all_research_items:
        c=it["category"]
        if c in group:
            group[c].append(it)
        else:
            group["其他"].append(it)
    for cat_name,qlist in group.items():
        if not qlist:continue
        doc.add_heading(f"▶ {cat_name}",level=2)
        for qi in qlist:
            doc.add_paragraph(f"【原题】{qi['question']}")
            for para in qi["research_content"].split("\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
            doc.add_paragraph("‑"*60)
    word_path = Path(OUT_RESEARCH_FOLDER)/"高考数学教研汇总.docx"
    doc.save(str(word_path))
    #导出pdf
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    c = canvas.Canvas(str(Path(OUT_RESEARCH_FOLDER)/"高考数学教研汇总.pdf"),pagesize=A4)
    c.setFont("STSong-Light",11)
    w,h = A4
    y = h-40
    def wline(txt):
        nonlocal y
        if y<40:
            c.showPage()
            c.setFont("STSong-Light",11)
            y=h-40
        c.drawString(30,y,txt[:110])
        y-=16
    wline("====================高考数学真题教研资料====================")
    for it in all_research_items:
        wline(f"【分类】{it['category']}")
        wline(f"【原题】{it['question'][:120]}")
        for line in it["research_content"].split("\n"):
            if line.strip():
                wline(line.strip())
        wline("="*70)
    c.save()
    print(f"\n✅教研完成！")
    print(f"Word：{word_path.resolve()}")
    print(f"PDF：{(Path(OUT_RESEARCH_FOLDER)/'高考数学教研汇总.pdf').resolve()}")

# ========================精细化内容生成（学科通用模板）========================
# 学科知识体系：每个学科一个专题列表（可扩展其他学科）
FINE_TOPICS = {
    "数学": {
        "第一篇 集合、逻辑与不等式基础": [
            "1. 集合的概念与表示", "2. 集合的运算", "3. 充分条件、必要条件与量词",
            "4. 不等式基础", "5. 基本不等式"],
        "第二篇 函数体系": [
            "6. 函数概念与定义域", "7. 函数单调性", "8. 函数奇偶性、周期性与对称性",
            "9. 二次函数", "10. 指数函数与对数函数", "11. 幂函数", "12. 函数图像变换",
            "13. 函数零点与方程", "14. 函数综合与参数"],
        "第三篇 三角函数与解三角形": [
            "15. 任意角与弧度制", "16. 三角函数定义与基本关系", "17. 诱导公式与恒等变换",
            "18. 三角函数图像与性质", "19. 解三角形"],
        "第四篇 平面向量与复数": ["20. 平面向量", "21. 复数"],
        "第五篇 数列": [
            "22. 数列概念与递推", "23. 等差数列", "24. 等比数列", "25. 数列求和", "26. 数列综合"],
        "第六篇 立体几何与空间向量": [
            "27. 空间几何体", "28. 空间点线面关系", "29. 空间平行与垂直证明", "30. 空间向量"],
        "第七篇 解析几何": [
            "31. 直线与斜率", "32. 圆", "33. 椭圆", "34. 双曲线", "35. 抛物线", "36. 圆锥曲线综合"],
        "第八篇 概率统计": [
            "37. 统计基础", "38. 抽样方法", "39. 古典概型与条件概率",
            "40. 全概率与贝叶斯思想", "41. 离散型随机变量", "42. 二项分布与正态分布"],
        "第九篇 计数原理与排列组合": [
            "43. 分类加法与分步乘法", "44. 排列、组合与二项式"],
    },
    "物理": {
        "第一篇 力学": [
            "运动的描述", "匀变速直线运动", "相互作用（重力、弹力、摩擦力）", "牛顿运动定律",
            "曲线运动与抛体", "圆周运动", "万有引力与航天", "功和功率", "动能定理", "机械能守恒定律"],
        "第二篇 电磁学": [
            "电场", "电势与电势能", "电容", "恒定电流", "磁场", "安培力与洛伦兹力",
            "电磁感应", "交变电流", "变压器与远距离输电"],
        "第三篇 热学与近代物理": [
            "分子动理论", "气体状态方程", "热力学定律", "动量与冲量", "动量守恒定律", "原子结构", "原子核"],
        "第四篇 实验与方法": [
            "力学实验", "电学实验", "读数与误差分析"],
    },
    "化学": {
        "第一篇 基本概念与理论": [
            "物质的量", "化学计量与阿伏伽德罗常数", "离子反应", "氧化还原反应",
            "元素周期律", "化学键", "化学反应速率", "化学平衡", "电离与水解", "原电池", "电解池"],
        "第二篇 元素及其化合物": [
            "钠及其化合物", "铝及其化合物", "铁及其化合物", "氯及其化合物", "硫及其化合物",
            "氮及其化合物", "硅及其化合物"],
        "第三篇 有机化学": [
            "烷烃烯烃炔烃", "苯及芳香烃", "卤代烃", "醇酚醛", "羧酸酯", "糖类油脂蛋白质",
            "有机合成与推断"],
        "第四篇 化学实验与计算": [
            "化学实验基本操作", "物质的检验与分离", "化学计算"],
    },
}

# 12 项精细化模板的固定结构（学科无关，供 LLM 生成时遵循）
FINE_TEMPLATE_SECTIONS = [
    "1. 核心概念与定义",
    "2. 必背公式/定理/定律（含常见变形与重要结论）",
    "3. 基本性质与规律",
    "4. A级基础题型（概念辨析、直接计算，2~3题，含完整题面与解析）",
    "5. B级常见题型（标准方法，4~5题，含完整题面与解析，覆盖本专题核心方法）",
    "6. C级综合题型（多知识点结合，3~4题，含完整题面与解析）",
    "7. D级拔高题型（以历年高考/会考真题题型为主，3~4题，含完整题面与解析，注明考法）",
    "8. 题型识别特征速查（看到XX→用XX，8条以上）",
    "9. 标准解题方法总流程",
    "10. 典型易错点诊断（5条以上）",
    "11. 重难点突破（含核心方法与技巧）",
    "12. 高考/会考真题映射与跨学科迁移",
]

FINE_PROMPT_TEMPLATE = r"""你是资深{subject}教师与教研专家。请为「{subject} · {topic}」专题编写一份完整、详实的精细化精讲讲义，严格按以下 12 个板块组织：

{template}

【数学符号书写规则 - 非常重要，必须严格遵守】
1. 一律使用 Unicode 纯文本数学符号，禁止使用 LaTeX 或 Markdown 数学标记：
   - 不要出现 \( \)、\[ \]、\frac、\sqrt、\geq、\leq、\cdot、\dots、\times、^、_、大括号 等 LaTeX 命令；
   - 分数线用斜杠：如 a+b 除以 2 写成 (a+b)/2；
   - 开方用 sqrt：如 √(x+1)、√3、∛；
   - 不等式用符号：≥、≤、≠、>、<；
   - 乘号用 × 或省略，不要用 *；
   - 上下标用普通写法：如 a²、b³、x₁、x₂、aⁿ、f'(x)、Sₙ；
   - 特殊符号直接写：π、±、∞、∈、⊆、∩、∪、⇒、⇔、∑、√、Δ。
2. 写出的文本必须能直接阅读，不要留下任何 LaTeX 代码。

【内容要求】
1. 内容准确、专业，面向高中{subject}复习与备考场景；
2. 题型部分每道题都必须给出【完整题面 + 分步解析 + 最终答案】，中档题（B/C级）要具体详细；
3. 拔高题（D级）以历年高考/会考真题题型为主，题目标注考法与年份（如“2023全国卷”），并给出贴近高考难度的完整解析；
4. 【必背公式/定理】板块必须补充本专题的重要变形、常见结论、推论与“秒杀结论”，例如基本不等式要包含：1的代换、凑定值、等号成立条件、调和≤几何≤算术≤平方不等式链、a²+b²≥2ab、ab≤((a+b)/2)²等；
5. 【重难点突破】必须写清本专题的核心方法与技巧（如基本不等式的“1的代换/常数代换、拆项、凑定值、换元、放缩”等），并配1~2个典型例子说明；
6. 【题型识别特征】至少8条；【典型易错点】至少5条；
7. 每个板块用「【板块名】」开头，直接输出正文，不要额外解释；
8. 总字数 6000~9000 字，保证内容丰富、可直接用于备课或学生复习。

【去AI味·教师口吻 - 非常重要，必须严格遵守】
1. 你是给真实高中生备课的老师，不是写报告的机器人。讲解要有老师讲课的口吻：直接、明确、有主次，像"这个考点常考，务必注意"这样的提醒；
2. 禁用套话连接词：不要出现"首先、其次、最后、此外、值得注意的是、不可否认、综上所述、总而言之、众所周知"这类机械过渡词，直接用内容推进；
3. 禁用强行洞察句式：不要用"不是A而是B""不仅是X更是Y""核心就一句话""归根结底"这类制造深刻感的表达；
4. 禁用空泛形容词和虚词：不要用"非常重要、非常关键、深刻的、全面的、系统性的"堆砌，把判断落到具体内容上（如"这里漏掉等号条件会失分"）；
5. 结尾不要强行升华：不要写"掌握这些就能……""数学是思维的体操"之类；板块结尾可以停在具体方法、易错提醒或下一板块的自然衔接；
6. 结构要真实：允许一句话短句、允许板块长短不一，不要每段都写成"观点+解释+例子+总结"的四段式；
7. 例题讲解要说清"为什么这样想"：先给思路（看到什么条件→想到什么方法），再给步骤，不要只机械罗列计算过程；
8. 允许自然的教学习惯用语：如"注意""这里容易错""考试时""这类题一般先……"，但不要过量堆砌；
9. 去AI味绝不牺牲准确性：公式、定理、术语、步骤必须完整准确，专业严谨始终优先。
"""


def _fine_llm(model_choice="auto"):
    """获取精细化生成用的 LLM。auto 默认用数学/快速模型（免费优先），reasoning 用推理模型"""
    llm_obj, label = get_llm_for_choice(model_choice, default="llm_fast")
    return llm_obj, label


FINE_CONTINUE_RULES = """【数学符号书写规则】一律使用 Unicode 纯文本数学符号，禁止使用 LaTeX/Markdown 数学标记（不要出现 \\( \\)、\\[ \\]、\\frac、\\sqrt、\\geq、\\cdot、^、_、花括号等命令）；分数线用斜杠如 (a+b)/2；开方用 √；不等式用 ≥、≤、≠；乘号用 ×；上下标用 a²、b₁、aⁿ、f'(x)。写出的文本必须能直接阅读。
【内容要求】题型部分每道题给出【完整题面 + 分步解析 + 最终答案】；中档题（B/C级）具体详细；拔高题（D级）以高考真题题型为主并标注考法与年份；重难点写清核心方法与技巧；题型识别至少8条；易错点至少5条。直接输出正文，不要额外解释。【字数要求 - 必须遵守】你本次负责的 4 个板块合计至少 2000 字，题型板块（含题目与解析）务必展开写满，宁可详实不可简略；每个板块都要有实质内容，禁止只用一行带过。
【去AI味·教师口吻】用老师讲课的口吻，直接、明确、有主次；禁用套话连接词（首先/其次/此外/值得注意的是/综上所述）；禁用"不是A而是B"等强行洞察句；不写空泛形容词堆砌；结尾不强行升华；例题先说"为什么这样想"再给步骤；允许自然教学习惯语（注意/这里容易错）；去AI味绝不牺牲公式定理与步骤的准确性。"""

def generate_fine_content(subject="数学", topic="函数单调性", model_choice="auto", progress_cb=None):
    """生成单个专题的精细化内容（Markdown 文本）。分段生成以突破单次输出长度限制，
    拼接时按板块编号去重，保证每个板块只出现一次。
    返回 (markdown文本, 模型显示名, 错误信息)"""
    import re as _re
    def cb(msg):
        if progress_cb:
            progress_cb(msg)

    llm_obj, label = _fine_llm(model_choice)

    # 分段：12个板块分3批，每批4个，逐批生成后拼接
    batches = [FINE_TEMPLATE_SECTIONS[i:i+4] for i in range(0, len(FINE_TEMPLATE_SECTIONS), 4)]

    # 段间引导：后一批衔接前一批，保持整体风格
    parts = []
    total = 0
    for bi, batch in enumerate(batches):
        template_lines = "\n".join(f"【{s}】" for s in batch)
        if bi == 0:
            prompt = FINE_PROMPT_TEMPLATE.format(subject=subject, topic=topic, template=template_lines)
        else:
            # 明确本次要生成的板块编号范围
            sec_nums = []
            for s in batch:
                m = _re.match(r"^(\d+)", s)
                if m:
                    sec_nums.append(m.group(1))
            nums_str = "、".join(sec_nums)
            prev_sections = "、".join(FINE_TEMPLATE_SECTIONS[:bi*4])
            prompt = (
                f"你正在编写「{subject} · {topic}」专题的精细化精讲讲义。"
                f"前面已经完成编号为 {prev_sections} 的板块。\n"
                f"现在请继续，本次【只输出】编号 {nums_str} 的以下板块（每个板块用【编号 名称】开头，不要重复前面已完成的板块）：\n\n"
                f"{template_lines}\n\n"
                + FINE_CONTINUE_RULES
            )
        cb(f"[{topic}] 正在生成第 {bi+1}/{len(batches)} 批（板块 {nums_str if bi>0 else '1-4'}）...")
        text = ""
        for attempt in range(2):
            try:
                resp = llm_obj.invoke(prompt)
                text = resp.content if hasattr(resp, "content") else str(resp)
                text = text.strip()
                if not text:
                    return "", label, f"第{bi+1}批返回为空"
                # 最低字数检查：每批4个板块至少1400字，不足则重试一次
                min_chars = 1400
                if len(text) < min_chars and attempt == 0:
                    cb(f"[{topic}] 第 {bi+1} 批仅 {len(text)} 字（低于{min_chars}），重新生成一次...")
                    prompt = prompt + f"\n\n【注意】你上一次输出太简略（仅 {len(text)} 字），请务必把本次的 4 个板块写充实：每道题给出完整题面+分步解析+最终答案，每个板块都要有实质内容，本次合计至少 2000 字。"
                    continue
                break
            except Exception as e:
                if attempt == 0:
                    cb(f"[{topic}] 第 {bi+1} 批调用失败，重试一次...")
                    continue
                return "", label, f"第{bi+1}批模型调用失败：{str(e)}"
        parts.append(text)
        total += len(text)
        cb(f"[{topic}] 第 {bi+1} 批完成（{len(text)} 字符）")

    # 拼接并按板块编号去重：保留每个编号第一次出现的板块内容
    def split_sections(text):
        """按 【n. 标题】 切分为 [(编号, 标题, 内容), ...]，非板块开头归入上一个板块"""
        secs = []
        cur_num, cur_title, cur_body = None, None, []
        for line in text.split("\n"):
            m = _re.match(r"^#*\s*【(\d+)[.、]?([^】]*)】", line)
            if m:
                if cur_num is not None:
                    secs.append((cur_num, cur_title, "\n".join(cur_body).strip()))
                cur_num = m.group(1)
                cur_title = m.group(2).strip()
                cur_body = []
            else:
                if cur_num is not None:
                    cur_body.append(line)
        if cur_num is not None:
            secs.append((cur_num, cur_title, "\n".join(cur_body).strip()))
        return secs

    seen = set()
    merged = []
    for part in parts:
        for num, title, body in split_sections(part):
            if num in seen:
                continue  # 板块已存在，跳过（去重）
            seen.add(num)
            merged.append(f"【{num}. {title}】\n{body}".rstrip())

    # 若没有按板块切分出内容（如部分批次未用【编号】开头），则原样拼接
    if not merged:
        md = "\n\n".join(parts)
    else:
        md = "\n\n".join(merged)

    cb(f"[{topic}] 全部生成完成，去重后共 {len(md)} 字符（{len(seen)} 个板块）")
    return md, label, None


def _latex_to_unicode(text):
    """将残留的 LaTeX/Markdown 数学标记转换为可读的 Unicode 数学符号"""
    import re as _re
    s = text
    # 去掉 LaTeX 内联/独立数学环境的包裹符
    s = s.replace("\\[", "").replace("\\]", "")
    s = s.replace("\\( ", "").replace("\\)", "").replace("\\) ", "")
    s = s.replace("\\(", "").replace("\\)", "")
    # 常见 LaTeX 命令 → Unicode
    s = s.replace("\\geq", "≥").replace("\\leq", "≤").replace("\\ne", "≠")
    s = s.replace("\\cdots", "⋯")
    s = s.replace("\\times", "×").replace("\\cdot", "·").replace("\\pm", "±")
    s = s.replace("\\infty", "∞").replace("\\pi", "π").replace("\\Delta", "Δ")
    s = s.replace("\\in", "∈").replace("\\subseteq", "⊆").replace("\\cap", "∩").replace("\\cup", "∪")
    s = s.replace("\\Rightarrow", "⇒").replace("\\Leftrightarrow", "⇔").replace("\\rightarrow", "→")
    s = s.replace("\\dots", "…").replace("\\ldots", "…")
    # \frac{a}{b} → (a)/(b)
    def _frac(m):
        a, b = m.group(1), m.group(2)
        a2 = a.replace("{", "").replace("}", "")
        b2 = b.replace("{", "").replace("}", "")
        return f"({a2})/({b2})"
    s = _re.sub(r"\\frac\{([^}]*)\}\{([^}]*)\}", _frac, s)
    # \sqrt[n]{x} → n次根号；\sqrt{x} → √(x)
    def _sqrt(m):
        if m.group(1):
            return f"∛({m.group(2)})" if m.group(1) == "3" else f"{m.group(1)}√({m.group(2)})"
        return f"√({m.group(2)})"
    s = _re.sub(r"\\sqrt(?:\[([^\]]*)\])?\{([^}]*)\}", _sqrt, s)
    # \bar{z} → z̄ ; \overline{z} → z̄
    s = _re.sub(r"\\bar\{([^}]*)\}", r"\1̄", s)
    s = _re.sub(r"\\overline\{([^}]*)\}", r"\1̄", s)
    # 去掉残留的反斜杠命令（未识别命令保留原字符但去掉 \）
    s = _re.sub(r"\\[a-zA-Z]+", "", s)
    # 下标数字/字母：a_1 → a₁, S_n → Sₙ
    sub_map = {"0":"₀","1":"₁","2":"₂","3":"₃","4":"₄","5":"₅","6":"₆","7":"₇","8":"₈","9":"₉","n":"ₙ","i":"ᵢ","k":"ₖ","m":"ₘ","j":"ⱼ","N":"ₙ"}
    def _sub(m):
        return "".join(sub_map.get(c, c) for c in m.group(1))
    s = _re.sub(r"_\{([^}]*)\}", _sub, s)
    s = _re.sub(r"_([0-9nNkmj])", lambda m: _sub(m), s)
    # 清理残留的花括号
    s = s.replace("{", "").replace("}", "")
    # 清理可能出现的 \ 和多余空格
    s = s.replace("\\", "")
    # 仅压缩连续空格，保留换行（保证 Word 分段）
    s = _re.sub(r"[ ]+", " ", s)
    s = _re.sub(r" *\n *", "\n", s)
    return s.strip()


def _fine_md_to_docx(md_text, subject, topic, out_path):
    """将精细化 Markdown 文本转换为 Word 文档（结构化，带标题层级，自动清理Markdown符号）"""
    md_text = _latex_to_unicode(md_text)
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    import re as _re

    doc = Document()
    sec = doc.sections
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.5)

    def set_font(run, size=12, bold=False, cn="宋体", en="Arial", color=None):
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = en
        r = run._element.rPr.rFonts
        r.set(qn("w:eastAsia"), cn)
        if color:
            run.font.color.rgb = RGBColor(*color)

    def add_center(text, size=18, bold=True, cn="黑体", after=18):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        set_font(p.add_run(text), size=size, bold=bold, cn=cn)

    def clean_md(s):
        """清理行首的 Markdown 标记"""
        s = s.replace("**", "").replace("__", "")  # 先删粗体标记（否则 * 会干扰列表判断）
        s = _re.sub(r"^#{1,6}\s*", "", s)      # ### → 去
        s = _re.sub(r"^[-*•]\s*", "• ", s)      # 列表符号统一为 •
        s = s.replace("`", "")
        s = _re.sub(r"^◆+", "◆ ", s)
        return s.strip()

    # 标题
    add_center(f"{subject} · {topic} · 精细化精讲")

    lines = md_text.split("\n")
    for line in lines:
        raw = line.strip()
        if not raw:
            continue
        s = clean_md(raw)

        # 板块标题：【n. xxx】或 ### 【n. xxx】→ 黑体15
        m = _re.match(r"^【(\d+[.、]?[^】]*)】\s*(.*)$", s)
        if m:
            sec_title = f"{m.group(1)} {m.group(2)}".strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            set_font(p.add_run(sec_title), size=15, bold=True, cn="黑体")
            continue

        # 子标题：题1/题2/例题1 或 1. xxx（作为小题标题）→ 宋体加粗13
        m2 = _re.match(r"^(题\d+[、.：: ]|例\d+[、.：: ]|【?\d+\.\s*[^。]{2,20}】?$)", s)
        is_subtitle = bool(m2) and len(s) < 40

        # 普通正文（含例题等）
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        if len(s) > 0 and not s.startswith(("①", "②", "③", "④", "⑤", "⑥", "-", "•", "（", "◆")):
            pf.first_line_indent = Pt(24)
        if is_subtitle:
            set_font(p.add_run(s), size=12, bold=True)
        else:
            set_font(p.add_run(s), size=12)

    doc.save(out_path)


def fine_generate_and_save(subject, topic, model_choice="auto", out_folder=None, progress_cb=None):
    """精细化生成并保存 Word。返回 (成功标志, 提示信息, 文件路径)"""
    if out_folder is None:
        out_folder = Path(OUT_RESEARCH_FOLDER) / "精细化内容"
    out_folder = Path(out_folder)
    out_folder.mkdir(parents=True, exist_ok=True)

    md_text, label, err = generate_fine_content(subject, topic, model_choice, progress_cb)
    if err:
        return False, f"❌ 生成失败：{err}", None

    safe = re.sub(r'[\\/:*?"<>|]', "_", topic)
    out_path = out_folder / f"{subject}_{safe}_精细化精讲.docx"
    try:
        _fine_md_to_docx(md_text, subject, topic, str(out_path))
    except Exception as e:
        return False, f"❌ Word 生成失败：{str(e)[:200]}", None

    return True, f"✅ 已生成：{out_path.name}（模型：{label}）", str(out_path)


# ========================主菜单========================
def print_main_menu():
    print("""
==================== 高考数学一体化Agent 主菜单 ====================
1 ｜单题交互模式（记忆+RAG+大题拆解+绘图+校验复核）
2 ｜编译/更新 LLM Wiki 知识库（Karpathy式，支持 txt/md/docx/pdf）
3 ｜批量做题（batch_input支持 txt / md / docx / pdf，输出docx+pdf）
4 ｜真题教研模块：批量读取raw_exam_files中PDF/Word，输出教研Word+PDF
q ｜退出程序
===================================================================
""")

if __name__ == "__main__":
    while True:
        print_main_menu()
        opt = input("请输入功能数字：").strip()
        if opt.lower()=="q":
            print("程序结束")
            break
        if opt=="1":
            thread_id = "chat_session01"
            cfg = {"configurable":{"thread_id":thread_id}, "recursion_limit": 10}
            print("\n====单题交互，clear清空记忆；exit退回主菜单====\n")
            while True:
                user_in = input("请输入题目：").strip()
                if user_in.lower()=="exit":
                    break
                if user_in.lower()=="clear":
                    graph_chat.update_state(cfg,{"messages":[]})
                    print("✅记忆清空")
                    continue
                t0 = time.time()
                ans_text = run_question(graph_chat, user_in, cfg)
                print(f"（本题用时 {time.time()-t0:.0f} 秒）")
                print("\n【输出】")
                print(ans_text)
                print("‑"*80)
        elif opt=="2":
            build_wiki()
        elif opt=="3":
            exts = (".txt",".md",".docx",".pdf")
            files = [f for f in os.listdir(BATCH_INPUT_FOLDER) if Path(f).suffix.lower() in exts]
            if not files:
                print(f"batch_input文件夹没有可识别试卷，支持格式：{exts}")
                continue
            for idx,f in enumerate(files):
                print(f"{idx}: {f}")
            sel = int(input("选择试卷序号："))
            target = files[sel]
            full_text,_ = batch_solve_file(target)
            pdf_out = os.path.join(OUTPUT_FOLDER, Path(target).stem+"_result.pdf")
            text_to_pdf(full_text, pdf_out)
        elif opt=="4":
            research_workflow()
        else:
            print("无效输入，请重新选择")
